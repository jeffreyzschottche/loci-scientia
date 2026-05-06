import base64
import binascii
import io
import logging
import math
import os
import re
import unicodedata
import uuid
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence
from zoneinfo import ZoneInfo
import httpx

from .devices_repo import DevicesRepository
from .knowledge_repository import KnowledgeRepository
from .schemas import ChatMessage, ChatRequest, Device, HistoryDocument, PromptDocument
from .settings import settings
from .translations import t
from .web_search import (
    WebSearchResult,
    format_results_for_prompt as _format_web_results_for_prompt,
)

logger = logging.getLogger(__name__)

MAX_CONTEXT_ITEMS = 3
MIN_CONTEXT_SCORE = 0.45
KNOWLEDGE_SNIPPET_LENGTH = 999999  # Characters to include from each knowledge chunk
PROMPT_TEMPLATE_PATH = Path(__file__).with_name("prompt.txt")
PROMPT_TEMPLATE_DIR = Path(__file__).with_name("prompt_templates")
PROMPT_LOG_PATH = Path(__file__).resolve().parents[2] / "promptlog.log"
API_PROMPT_LOG_PATH = Path(__file__).resolve().parents[2] / "apiprompt.log"
THINKING_TAG_RE = re.compile(r"<think>(.*?)</think>", re.IGNORECASE | re.DOTALL)
_TOKEN_PATTERN = re.compile(r"\S+")
MIN_IMAGE_DIMENSION = 32
MAX_DOCUMENTS_PER_REQUEST = 3
MAX_DOCUMENT_BYTES = 5 * 1024 * 1024
MAX_TOTAL_DOCUMENT_BYTES = 12 * 1024 * 1024
MAX_DOCUMENT_CHARS = 40_000
MAX_TOTAL_DOCUMENT_CHARS = 100_000
MAX_HISTORY_DOCUMENT_CHARS = 12_000
MAX_HISTORY_DOCUMENT_CHARS_PER_DOC = 4_000
MAX_HISTORY_IMAGES = 8
THINKING_MODEL_MARKERS = (
    "qwen3",
    "qwen 3",
    "qwen3.5",
    "deepseek-r1",
    "deepseek_r1",
    "deepseek-v3.1",
    "gpt-oss",
)
_KNOWN_BAD_OUTPUT_FRAGMENTS = (
    "receptor model ad aanpassingen",
    "hub heeft geen tools toegewezen",
    "ik kan geen tools gebruiken",
)


@dataclass
class ApiLogEntry:
    """Structured log entry for API conversations."""
    conversation_id: str
    request_id: str
    timestamp: str
    source: str  # "local" or "api"
    endpoint: str  # "/api/v1/ask" or "/api/v1/ask/stream"
    mode: Optional[str] = None
    thinking_enabled: Optional[bool] = None
    user_name: Optional[str] = None
    device_id: Optional[str] = None
    token_prefix: Optional[str] = None
    original_prompt: str = ""
    new_chat: bool = False  # True if this request cleared the history
    history_length: int = 0
    images_count: int = 0
    documents_count: int = 0
    context_devices: List[Dict[str, Any]] = field(default_factory=list)
    context_knowledge: List[Dict[str, Any]] = field(default_factory=list)
    context_documents: List[Dict[str, Any]] = field(default_factory=list)
    final_prompt_length: int = 0
    response_preview: str = ""
    error: Optional[str] = None


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    text: str
    source_bytes: int
    truncated: bool = False


def generate_request_id() -> str:
    """Generate a unique request ID."""
    return str(uuid.uuid4())[:8]


def generate_conversation_id(token: str) -> str:
    """Generate a conversation ID based on token (first 8 chars of hash)."""
    import hashlib
    return hashlib.sha256(token.encode()).hexdigest()[:8]


def log_api_request(entry: ApiLogEntry) -> None:
    """Write structured API log entry to apiprompt.log."""
    try:
        API_PROMPT_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)

        separator = "=" * 80
        new_chat_marker = " [NEW CHAT]" if entry.new_chat else ""
        lines = [
            separator,
            f"[{entry.timestamp}]{new_chat_marker}",
            f"Conversation ID: {entry.conversation_id}",
            f"Request ID:      {entry.request_id}",
            f"Source:          {entry.source}",
            f"Endpoint:        {entry.endpoint}",
            f"Mode:            {entry.mode or 'default'}",
            f"Thinking:        {'on' if entry.thinking_enabled else 'off' if entry.thinking_enabled is not None else 'auto'}",
            f"User:            {entry.user_name or 'unknown'}",
            f"Device ID:       {entry.device_id or 'unknown'}",
            f"Token:           {entry.token_prefix or 'unknown'}...",
            f"New Chat:        {entry.new_chat}",
            "",
            f"--- Original Prompt ({len(entry.original_prompt)} chars) ---",
            entry.original_prompt[:500] + ("..." if len(entry.original_prompt) > 500 else ""),
            "",
            f"History messages: {entry.history_length}",
            f"Images attached:  {entry.images_count}",
            f"Documents:        {entry.documents_count}",
            "",
            "--- Context Found ---",
            f"Devices ({len(entry.context_devices)}):",
        ]

        if entry.context_devices:
            for d in entry.context_devices:
                lines.append(f"  - {d.get('user_name', '?')} / {d.get('device_name', '?')} (score: {d.get('score', 0):.3f})")
        else:
            lines.append("  (none)")

        lines.append(f"Knowledge ({len(entry.context_knowledge)}):")
        if entry.context_knowledge:
            for k in entry.context_knowledge:
                lines.append(f"  - [{k.get('doc_id', '?')}] {k.get('title', '?')} (score: {k.get('score', 0):.3f})")
                lines.append(f"    {k.get('snippet', '')}")
        else:
            lines.append("  (none)")

        lines.append(f"Documents ({len(entry.context_documents)}):")
        if entry.context_documents:
            for d in entry.context_documents:
                lines.append(
                    f"  - {d.get('filename', '?')} [{d.get('file_type', '?')}] "
                    f"{d.get('chars', 0)} chars / {d.get('bytes', 0)} bytes"
                )
        else:
            lines.append("  (none)")

        lines.extend([
            "",
            f"Final prompt length: {entry.final_prompt_length} chars",
        ])

        if entry.response_preview:
            lines.extend([
                "",
                "--- Response Preview ---",
                entry.response_preview[:300] + ("..." if len(entry.response_preview) > 300 else ""),
            ])

        if entry.error:
            lines.extend([
                "",
                f"--- ERROR ---",
                entry.error,
            ])

        lines.append("")

        with API_PROMPT_LOG_PATH.open("a", encoding="utf-8") as handle:
            handle.write("\n".join(lines) + "\n")

    except OSError as exc:
        logger.warning("Kon API log niet schrijven: %s", exc)

devices_repo = DevicesRepository()
knowledge_repo = KnowledgeRepository()


def describe_device(device: Device) -> str:
    lines = [
        f"{device.user_name} ({device.device_name})",
        f"  ✉ {device.email}" if device.email else "",
        f"  ☎ {device.phone}" if device.phone else "",
    ]
    return "\n".join(line for line in lines if line)


def _flatten_multiline(text: str) -> str:
    return " ".join(line.strip() for line in text.splitlines() if line.strip())


def _clean_document_text(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").splitlines())


def _strip_data_url(value: str) -> str:
    raw = (value or "").strip()
    if not raw.startswith("data:"):
        return raw
    _, _, b64 = raw.partition("base64,")
    return b64.strip() if b64 else ""


def _decode_base64_image(value: str, index: int) -> bytes:
    raw = _strip_data_url(value)
    if not raw:
        raise ValueError(f"Image {index} has no base64 payload.")
    try:
        return base64.b64decode(raw, validate=True)
    except binascii.Error as exc:
        raise ValueError(f"Image {index} contains invalid base64 data.") from exc


def _validate_image_payload(value: str, index: int) -> None:
    source = _decode_base64_image(value, index)
    try:
        from PIL import Image, UnidentifiedImageError
    except Exception:
        return

    try:
        with Image.open(io.BytesIO(source)) as image:
            width, height = image.size
            image_format = (image.format or "unknown").upper()
    except UnidentifiedImageError as exc:
        raise ValueError(f"Image {index} is not a supported image format.") from exc

    if width < MIN_IMAGE_DIMENSION or height < MIN_IMAGE_DIMENSION:
        raise ValueError(
            f"Image {index} is too small for Ollama vision ({width}x{height}). "
            f"Minimum size is {MIN_IMAGE_DIMENSION}x{MIN_IMAGE_DIMENSION}."
        )

    logger.info(
        "Validated image %s for Ollama: format=%s size=%sx%s bytes=%s",
        index,
        image_format,
        width,
        height,
        len(source),
    )


def _detect_document_type(document: PromptDocument) -> str:
    file_name = (document.filename or "").strip().lower()
    if "." in file_name:
        ext = file_name.rsplit(".", 1)[-1]
        if ext in {"pdf", "docx", "xls", "xlsx", "txt"}:
            return ext
    content_type = (document.content_type or "").strip().lower()
    type_map = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "docx",
        "application/vnd.ms-excel": "xls",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "text/plain": "txt",
    }
    for key, value in type_map.items():
        if key in content_type:
            return value
    return ""


def _decode_base64_document(document: PromptDocument) -> bytes:
    raw = _strip_data_url(document.data)
    if not raw:
        raise ValueError(f"Document '{document.filename}' has no base64 payload.")
    try:
        return base64.b64decode(raw, validate=True)
    except binascii.Error as exc:
        raise ValueError(f"Document '{document.filename}' contains invalid base64 data.") from exc


def _extract_pdf_text(source: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - dependency absence
        raise ValueError("PDF parsing requires 'pypdf'. Install backend dependencies first.") from exc

    reader = PdfReader(io.BytesIO(source))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx_text(source: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency absence
        raise ValueError("DOCX parsing requires 'python-docx'. Install backend dependencies first.") from exc

    document = Document(io.BytesIO(source))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [str(cell.text).strip() for cell in row.cells if str(cell.text).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx_text(source: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover - dependency absence
        raise ValueError("XLSX parsing requires 'openpyxl'. Install backend dependencies first.") from exc

    workbook = load_workbook(filename=io.BytesIO(source), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell not in (None, "")]
            if cells:
                parts.append(" | ".join(cells))
    workbook.close()
    return "\n".join(parts)


def _extract_xls_text(source: bytes) -> str:
    try:
        import xlrd
    except Exception as exc:  # pragma: no cover - dependency absence
        raise ValueError("XLS parsing requires 'xlrd'. Install backend dependencies first.") from exc

    workbook = xlrd.open_workbook(file_contents=source, on_demand=True)
    parts: list[str] = []
    for sheet in workbook.sheets():
        parts.append(f"Sheet: {sheet.name}")
        for row_idx in range(sheet.nrows):
            cells: list[str] = []
            for value in sheet.row_values(row_idx):
                cell_text = str(value).strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))
    workbook.release_resources()
    return "\n".join(parts)


def _extract_txt_text(source: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            return source.decode(encoding)
        except UnicodeDecodeError:
            continue
    return source.decode("utf-8", errors="replace")


def _extract_document_text(file_type: str, source: bytes) -> str:
    if file_type == "pdf":
        return _extract_pdf_text(source)
    if file_type == "docx":
        return _extract_docx_text(source)
    if file_type == "xlsx":
        return _extract_xlsx_text(source)
    if file_type == "xls":
        return _extract_xls_text(source)
    if file_type == "txt":
        return _extract_txt_text(source)
    raise ValueError(f"Unsupported document type: {file_type}")


def normalize_documents(documents: Optional[Sequence[PromptDocument]]) -> list[ParsedDocument]:
    if not documents:
        return []
    if len(documents) > MAX_DOCUMENTS_PER_REQUEST:
        raise ValueError(
            f"Too many documents in one request ({len(documents)}). "
            f"Maximum allowed is {MAX_DOCUMENTS_PER_REQUEST}."
        )

    parsed_documents: list[ParsedDocument] = []
    total_bytes = 0
    total_chars = 0

    for index, document in enumerate(documents, 1):
        filename = (document.filename or "").strip() or f"document-{index}"
        file_type = _detect_document_type(document)
        if not file_type:
            raise ValueError(
                f"Document '{filename}' is not supported. Allowed: pdf, docx, xls, xlsx, txt."
            )

        source = _decode_base64_document(document)
        source_bytes = len(source)
        if source_bytes > MAX_DOCUMENT_BYTES:
            raise ValueError(
                f"Document '{filename}' is too large ({source_bytes} bytes). "
                f"Maximum per document is {MAX_DOCUMENT_BYTES} bytes."
            )

        total_bytes += source_bytes
        if total_bytes > MAX_TOTAL_DOCUMENT_BYTES:
            raise ValueError(
                f"Total document payload is too large ({total_bytes} bytes). "
                f"Maximum per request is {MAX_TOTAL_DOCUMENT_BYTES} bytes."
            )

        extracted = _clean_document_text(_extract_document_text(file_type, source)).strip()
        if not extracted:
            continue

        truncated = False
        if len(extracted) > MAX_DOCUMENT_CHARS:
            extracted = extracted[:MAX_DOCUMENT_CHARS]
            truncated = True

        remaining_chars = MAX_TOTAL_DOCUMENT_CHARS - total_chars
        if remaining_chars <= 0:
            break
        if len(extracted) > remaining_chars:
            extracted = extracted[:remaining_chars]
            truncated = True

        total_chars += len(extracted)
        parsed_documents.append(
            ParsedDocument(
                filename=filename,
                file_type=file_type,
                text=extracted,
                source_bytes=source_bytes,
                truncated=truncated,
            )
        )

    return parsed_documents


def _format_document_context_lines(documents: Optional[Sequence[ParsedDocument]]) -> list[str]:
    if not documents:
        return []
    lines: list[str] = []
    for index, document in enumerate(documents, 1):
        lines.append(f"- Document {index}: {document.filename} [{document.file_type}]")
        kenmerken = f"{len(document.text)} tekens uit {document.source_bytes} bytes"
        if document.truncated:
            kenmerken = f"{kenmerken}; ingekort"
        lines.append(f"  Kenmerken: {kenmerken}.")
        lines.append("  Geextracteerde inhoud:")
        for text_line in document.text.splitlines() or [""]:
            lines.append(f"  {text_line}")
    return lines


def history_documents_from_parsed(
    documents: Optional[Sequence[ParsedDocument]],
) -> list[HistoryDocument]:
    if not documents:
        return []

    history_documents: list[HistoryDocument] = []
    remaining_chars = MAX_HISTORY_DOCUMENT_CHARS

    for document in documents:
        if remaining_chars <= 0:
            break
        text = (document.text or "").strip()
        if not text:
            continue
        limit = min(MAX_HISTORY_DOCUMENT_CHARS_PER_DOC, remaining_chars)
        truncated = document.truncated
        if len(text) > limit:
            text = text[:limit]
            truncated = True
        remaining_chars -= len(text)
        history_documents.append(
            HistoryDocument(
                filename=document.filename,
                file_type=document.file_type,
                text=text,
                source_bytes=document.source_bytes,
                truncated=truncated,
            )
        )

    return history_documents


def _document_details(documents: Optional[Sequence[ParsedDocument]]) -> List[Dict[str, Any]]:
    if not documents:
        return []
    details: List[Dict[str, Any]] = []
    for document in documents:
        details.append(
            {
                "filename": document.filename,
                "file_type": document.file_type,
                "bytes": document.source_bytes,
                "chars": len(document.text),
                "truncated": document.truncated,
            }
        )
    return details


def _format_history_document_lines(documents: Optional[Sequence[HistoryDocument]]) -> list[str]:
    if not documents:
        return []
    lines: list[str] = []
    lines.append(f"[Documenten: {len(documents)}]")
    for index, document in enumerate(documents, 1):
        kenmerken = f"{len(document.text)} tekens uit {document.source_bytes} bytes"
        if document.truncated:
            kenmerken = f"{kenmerken}; ingekort"
        lines.append(f"- Document {index}: {document.filename} [{document.file_type}]")
        lines.append(f"  Kenmerken: {kenmerken}.")
        lines.append("  Geextracteerde inhoud:")
        for text_line in document.text.splitlines() or [""]:
            lines.append(f"  {text_line}")
    return lines


def _format_attachment_context_lines(
    *,
    images_count: int = 0,
    documents: Optional[Sequence[ParsedDocument]] = None,
) -> list[str]:
    if images_count <= 0 and not documents:
        return []

    lines = ["> bijlagen:"]
    if images_count > 0:
        noun = "afbeelding" if images_count == 1 else "afbeeldingen"
        verb = "is" if images_count == 1 else "zijn"
        lines.append(
            f"- Er {verb} {images_count} {noun} meegestuurd. "
            "De visuele inhoud is apart beschikbaar voor het model; baseer visuele claims alleen op wat echt zichtbaar is."
        )

    lines.extend(_format_document_context_lines(documents))
    return lines


def _format_pages(pages: Any) -> str:
    if isinstance(pages, (list, tuple)):
        numbers = [str(p) for p in pages if isinstance(p, int) or (isinstance(p, str) and p.strip())]
        if not numbers:
            return ""
        if len(numbers) == 1:
            return f"p. {numbers[0]}"
        return f"p. {numbers[0]}-{numbers[-1]}"
    if isinstance(pages, int):
        return f"p. {pages}"
    return ""


def _knowledge_source_label(payload: Dict[str, Any]) -> str:
    """Human-readable source: 'Filename — p. 3' or 'Title — p. 3' fallback."""
    filename = str(payload.get("original_filename") or "").strip()
    title = str(payload.get("document_title") or payload.get("doc_id") or "doc").strip()
    label = filename or title
    pages_label = _format_pages(payload.get("pages"))
    return f"{label} — {pages_label}" if pages_label else label


def _build_citation(payload: Dict[str, Any], score: float) -> Dict[str, Any]:
    snippet_text = _flatten_multiline(str(payload.get("text") or ""))
    snippet = snippet_text[:240] + ("..." if len(snippet_text) > 240 else "")
    return {
        "doc_id": payload.get("doc_id"),
        "chunk_id": payload.get("chunk_id"),
        "title": payload.get("document_title") or payload.get("doc_id"),
        "original_filename": payload.get("original_filename"),
        "pages": payload.get("pages"),
        "section": payload.get("section"),
        "score": round(float(score or 0.0), 4),
        "snippet": snippet,
    }


def _gather_context_with_details(prompt_text: str) -> tuple[list[str], Dict[str, List]]:
    """Gather context lines AND detailed info for logging + citations."""
    scored: list[tuple[str, str, float]] = []
    details: Dict[str, List] = {
        "devices": [],
        "knowledge": [],
        "citations": [],
    }

    device_hits = devices_repo.search_devices(prompt_text, limit=5)
    for device, score in device_hits:
        scored.append(
            (
                "device",
                _flatten_multiline(describe_device(device)),
                float(score or 0.0),
            )
        )
        details["devices"].append({
            "user_name": device.user_name,
            "device_name": device.device_name,
            "score": float(score or 0.0),
        })

    knowledge_hits = knowledge_repo.search_chunks(
        prompt_text,
        limit=MAX_CONTEXT_ITEMS,
        score_threshold=MIN_CONTEXT_SCORE,
    )
    for payload, score in knowledge_hits:
        text = _flatten_multiline(str(payload.get("text") or ""))
        if not text:
            continue
        source_label = _knowledge_source_label(payload)
        snippet = text[:KNOWLEDGE_SNIPPET_LENGTH] + ("..." if len(text) > KNOWLEDGE_SNIPPET_LENGTH else "")
        desc = f"[{source_label}] {snippet}"
        scored.append(("knowledge", desc, float(score or 0.0)))
        details["knowledge"].append({
            "doc_id": payload.get("doc_id"),
            "title": payload.get("document_title") or payload.get("doc_id"),
            "original_filename": payload.get("original_filename"),
            "pages": payload.get("pages"),
            "score": float(score or 0.0),
            "snippet": text[:100] + ("..." if len(text) > 100 else ""),
        })
        details["citations"].append(_build_citation(payload, score))

    if not scored:
        return [], details

    scored.sort(key=lambda item: item[2], reverse=True)
    limited = scored[:MAX_CONTEXT_ITEMS]
    lines: list[str] = []
    for idx, (kind, desc, score) in enumerate(limited, 1):
        lines.append(f"{idx}. [{kind}] score {score:.3f}: {desc}")
    return lines, details


def _normalize_mode(mode: Optional[str]) -> str:
    return (mode or "").strip().lower()


def _mode_filename(mode: str) -> str:
    slug = "".join(ch.lower() if ch.isalnum() else "_" for ch in mode.strip()).strip("_")
    return f"{slug}.txt" if slug else ""


def _resolve_mode(mode: Optional[str]) -> Optional[str]:
    normalized = _normalize_mode(mode)
    if not normalized:
        return None
    for configured in settings.prompt_modes:
        candidate = configured.strip()
        if candidate and candidate.lower() == normalized:
            return candidate
    return None


def _read_template(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8").strip()
    except FileNotFoundError:
        return ""


def _prompt_template(mode: Optional[str] = None) -> str:
    """Resolve prompt template for the selected mode with a default fallback."""
    resolved_mode = _resolve_mode(mode)
    if resolved_mode:
        filename = _mode_filename(resolved_mode)
        if filename:
            mode_template = _read_template(PROMPT_TEMPLATE_DIR / filename)
            if mode_template:
                return mode_template

    default_template = _read_template(PROMPT_TEMPLATE_PATH)
    if default_template:
        return default_template
    return t("prompt_default_template")


def _read_env_value(key: str) -> Optional[str]:
    env_value = os.environ.get(key)
    if env_value:
        return env_value.strip()
    env_path = Path(__file__).resolve().parents[2] / ".env"
    try:
        lines = env_path.read_text(encoding="utf-8").splitlines()
    except FileNotFoundError:
        return None
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in line:
            continue
        name, value = line.split("=", 1)
        if name.strip() != key:
            continue
        value = value.strip()
        if (value.startswith('"') and value.endswith('"')) or (
            value.startswith("'") and value.endswith("'")
        ):
            value = value[1:-1]
        return value.strip()
    return None


def _resolve_timezone_name() -> str:
    return _read_env_value("TIMEZONE") or "UTC"


def _resolve_today_in_timezone(timezone_name: str) -> datetime:
    try:
        tzinfo = ZoneInfo(timezone_name)
    except Exception:
        tzinfo = timezone.utc
        timezone_name = "UTC"
    configured_today = _read_env_value("AITJE_TODAY_DATE")
    if configured_today:
        try:
            parsed_date = datetime.fromisoformat(configured_today)
            return parsed_date.replace(tzinfo=tzinfo)
        except ValueError:
            pass
    return datetime.now(tzinfo)


def _resolve_language_label(value: Optional[str]) -> str:
    normalized = (value or "").strip().lower()
    if normalized.startswith("en"):
        return "English"
    return "Nederlands"


def _system_prompt_lines() -> list[str]:
    timezone_name = _resolve_timezone_name()
    today = _resolve_today_in_timezone(timezone_name)
    response_language = _resolve_language_label(_read_env_value("AITJE_RESPONSE_LANGUAGE"))
    application_language = _resolve_language_label(_read_env_value("LANGUAGE"))
    weekday_name = today.strftime("%A")
    return [
        "> system configuration:",
        f"- Application language: {application_language}",
        f"- Preferred response language: {response_language}",
        f"- Timezone: {timezone_name}",
        f"- Configured current date: {today.strftime('%Y-%m-%d')}",
        f"- Configured current weekday: {weekday_name}",
        "- Use the preferred response language by default unless the user explicitly asks for another language.",
    ]


def _format_history_lines(history: Optional[Sequence[ChatMessage]]) -> list[str]:
    if not history:
        return []
    lines: list[str] = []
    for idx, message in enumerate(history, 1):
        role = (message.role or "").lower()
        if role == "assistant":
            speaker = t("speaker_assistant")
        elif role == "system":
            speaker = t("speaker_system")
        else:
            speaker = t("speaker_user")
        content_parts: list[str] = []
        content = _flatten_multiline(message.content or "")
        if content:
            content_parts.append(content)
        images_count = len(getattr(message, "images", []) or [])
        if images_count > 0:
            marker = f"[Afbeeldingen: {images_count}]"
            if marker not in content_parts:
                content_parts.append(marker)
        document_lines = _format_history_document_lines(getattr(message, "documents", []) or [])
        if document_lines:
            content_parts.extend(document_lines)
        if not content_parts:
            continue
        rendered = "\n".join(content_parts)
        lines.append(f"{idx}. {speaker}: {rendered}")
    return lines


def collect_prompt_images(
    history: Optional[Sequence[ChatMessage]] = None,
    current_images: Optional[Sequence[str]] = None,
) -> list[str]:
    previous_images: list[str] = []
    seen: set[str] = set()

    for message in history or []:
        for image in getattr(message, "images", []) or []:
            data = (image or "").strip()
            if not data or data in seen:
                continue
            seen.add(data)
            previous_images.append(data)

    if len(previous_images) > MAX_HISTORY_IMAGES:
        previous_images = previous_images[-MAX_HISTORY_IMAGES:]
        seen = set(previous_images)

    combined = list(previous_images)
    for image in current_images or []:
        data = (image or "").strip()
        if not data or data in seen:
            continue
        seen.add(data)
        combined.append(data)
    return combined




def build_augmented_prompt(
    user_prompt: str,
    history: Optional[Sequence[ChatMessage]] = None,
    images_count: int = 0,
    documents: Optional[Sequence[ParsedDocument]] = None,
    mode: Optional[str] = None,
    web_results: Optional[Sequence[WebSearchResult]] = None,
) -> str:
    prompt, _ = build_augmented_prompt_with_details(
        user_prompt,
        history=history,
        images_count=images_count,
        documents=documents,
        mode=mode,
        web_results=web_results,
    )
    return prompt


def build_augmented_prompt_with_details(
    user_prompt: str,
    history: Optional[Sequence[ChatMessage]] = None,
    images_count: int = 0,
    documents: Optional[Sequence[ParsedDocument]] = None,
    mode: Optional[str] = None,
    web_results: Optional[Sequence[WebSearchResult]] = None,
) -> tuple[str, Dict[str, List]]:
    """Build augmented prompt AND return context details for logging."""
    base_lines: list[str] = []
    template = _prompt_template(mode=mode)
    if template:
        base_lines.append(template)
        base_lines.append("")

    base_lines.extend(_system_prompt_lines())
    base_lines.append("")

    history_lines = _format_history_lines(history)
    if history_lines:
        base_lines.append(t("previous_chat"))
        base_lines.extend(history_lines)
        base_lines.append("")

    attachment_lines = _format_attachment_context_lines(
        images_count=images_count,
        documents=documents,
    )
    if attachment_lines:
        base_lines.extend(attachment_lines)
        base_lines.append("")

    context_lines, context_details = _gather_context_with_details(user_prompt)
    context_details["documents"] = _document_details(documents)
    if context_lines:
        base_lines.append("> context:")
        base_lines.extend(context_lines)
        base_lines.append("")

    if web_results:
        web_lines = _format_web_results_for_prompt(list(web_results))
        if web_lines:
            base_lines.extend(web_lines)
            base_lines.append(
                "Gebruik bovenstaande webbronnen waar relevant en verwijs in het antwoord"
                " naar de bron met titel of URL. Als de bronnen geen antwoord geven, zeg dat dan."
            )
            base_lines.append("")

    base_lines.append(f"{t('current_question')} {user_prompt.strip()}")
    return "\n".join(base_lines).strip(), context_details


def prepare_prompt(
    req: ChatRequest,
    history: Optional[Sequence[ChatMessage]] = None,
    documents: Optional[Sequence[ParsedDocument]] = None,
    web_results: Optional[Sequence[WebSearchResult]] = None,
) -> tuple[str, list[str]]:
    """Normalize request payload and return the prompt used for context checks."""
    history_to_use = req.history if history is None else history
    normalized_images = normalize_images(req.images)
    prompt = build_augmented_prompt(
        req.prompt,
        history=history_to_use,
        images_count=len(normalized_images),
        documents=documents,
        mode=req.mode,
        web_results=web_results,
    )
    return prompt, normalized_images


def estimate_prompt_tokens(text: str) -> int:
    if not text:
        return 0
    wordish = len(_TOKEN_PATTERN.findall(text))
    char_est = math.ceil(len(text) / 4)
    byte_est = math.ceil(len(text.encode("utf-8")) / 4)
    # Heuristic: use the largest estimate to avoid undercounting.
    return max(wordish, char_est, byte_est)


def log_prompt(final_prompt: str) -> None:
    try:
        PROMPT_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        with PROMPT_LOG_PATH.open("a", encoding="utf-8") as handle:
            handle.write(f"[{datetime.utcnow().isoformat()}Z]\n{final_prompt}\n\n")
    except OSError:
        pass


def _build_summary_prompt(
    history_lines: Sequence[str],
    existing_summary: Optional[str] = None,
) -> str:
    lines: list[str] = [
        t("summary_instruction_1"),
        t("summary_instruction_2"),
        t("summary_instruction_3"),
        t("summary_instruction_4"),
        "",
    ]
    if existing_summary:
        lines.append(t("existing_summary"))
        lines.append(existing_summary.strip())
        lines.append("")
    lines.append(t("new_messages"))
    lines.extend(history_lines)
    return "\n".join(lines).strip()


def _fallback_summary(existing_summary: Optional[str], history_lines: Sequence[str]) -> str:
    combined = []
    if existing_summary:
        combined.append(existing_summary.strip())
    combined.append(" ".join(history_lines))
    text = " ".join(part for part in combined if part)
    return text[:1200].rstrip()


async def summarize_history(
    history: Optional[Sequence[ChatMessage]] = None,
    existing_summary: Optional[str] = None,
) -> str:
    history_lines = _format_history_lines(history)
    if not history_lines:
        return (existing_summary or "").strip()
    prompt = _build_summary_prompt(history_lines, existing_summary)
    try:
        summary = await _call_ollama(prompt, options={"num_predict": 256})
    except Exception as exc:  # pragma: no cover - netwerkfout
        logger.warning("%s %s", t("summary_failed"), exc)
        summary = _fallback_summary(existing_summary, history_lines)
    return (summary or "").strip()


def normalize_images(images: Optional[Sequence[str]]) -> list[str]:
    if not images:
        return []
    normalized: list[str] = []
    for index, item in enumerate(images, 1):
        if not item:
            continue
        data = item.strip()
        if not data:
            continue
        if data.startswith("data:"):
            _, _, b64 = data.partition("base64,")
            if not b64:
                raise ValueError(f"Image {index} has an invalid data URL without base64 payload.")
            data = b64.strip()
        _validate_image_payload(data, index)
        normalized.append(data)
    return normalized


def model_supports_ollama_thinking(model: str) -> bool:
    normalized = (model or "").strip().lower()
    return any(marker in normalized for marker in THINKING_MODEL_MARKERS)


def build_ollama_generate_payload(
    prompt: str,
    *,
    stream: bool,
    images: Optional[Sequence[str]] = None,
    thinking: Optional[bool] = None,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "model": settings.ollama_model,
        "prompt": prompt,
        "stream": stream,
    }
    if images:
        payload["images"] = list(images)
    max_context = settings.ollama_max_context.get(settings.ollama_model)
    if isinstance(max_context, int) and max_context > 0:
        payload["options"] = {"num_ctx": max_context}
    if model_supports_ollama_thinking(settings.ollama_model):
        payload["think"] = True if thinking is None else bool(thinking)
    return payload


def split_ollama_response_parts(
    response_text: Optional[str],
    thinking_text: Optional[str] = None,
) -> tuple[str, str]:
    response = (response_text or "").strip()
    thinking_parts = [part.strip() for part in [(thinking_text or "").strip()] if part.strip()]
    if response:
        inline_thinking = [match.strip() for match in THINKING_TAG_RE.findall(response) if match.strip()]
        if inline_thinking:
            thinking_parts.extend(inline_thinking)
            response = THINKING_TAG_RE.sub("", response).strip()
    seen: set[str] = set()
    ordered_thinking: list[str] = []
    for part in thinking_parts:
        if part and part not in seen:
            seen.add(part)
            ordered_thinking.append(part)
    return ("\n\n".join(ordered_thinking).strip(), response)


def _fallback_invalid_generation_response() -> str:
    return (
        "Het model gaf een onbruikbaar antwoord terug. "
        "Dit wijst meestal op een instabiele model/runtime-combinatie. "
        f"Probeer een ander model of schakel Vulkan uit en probeer opnieuw "
        f"(huidig model: '{settings.ollama_model}')."
    )


def invalid_model_output_message() -> str:
    return _fallback_invalid_generation_response()


def _looks_like_degenerate_repetition(text: str) -> bool:
    tokens = [token.lower() for token in _TOKEN_PATTERN.findall(text)]
    if len(tokens) < 5:
        return False
    counts = Counter(tokens)
    unique_count = len(counts)
    most_common_count = counts.most_common(1)[0][1]
    if unique_count <= 2 and most_common_count >= max(4, math.ceil(len(tokens) * 0.6)):
        return True
    return False


def _looks_like_garbled_text(text: str) -> bool:
    compact = (text or "").strip()
    if len(compact) < 24:
        return False

    non_space_chars = [char for char in compact if not char.isspace()]
    if not non_space_chars:
        return False

    letter_count = sum(1 for char in non_space_chars if char.isalpha())
    replacement_count = sum(1 for char in non_space_chars if char == "\ufffd")
    control_like_count = sum(
        1
        for char in non_space_chars
        if unicodedata.category(char) in {"Cc", "Cf", "Cs", "Co", "Cn"}
    )
    word_count = len(re.findall(r"\w+", compact, flags=re.UNICODE))
    space_ratio = sum(1 for char in compact if char.isspace()) / max(len(compact), 1)
    letter_ratio = letter_count / len(non_space_chars)
    corrupted_ratio = (replacement_count + control_like_count) / len(non_space_chars)

    if replacement_count >= 2:
        return True
    if corrupted_ratio >= 0.05:
        return True
    if word_count <= 2 and letter_ratio < 0.55 and len(compact) >= 32:
        return True
    if space_ratio < 0.04 and letter_ratio < 0.65 and len(compact) >= 48:
        return True
    return False


def is_invalid_model_output(text: Optional[str]) -> bool:
    normalized = " ".join((text or "").strip().lower().split())
    if not normalized:
        return False
    if any(fragment in normalized for fragment in _KNOWN_BAD_OUTPUT_FRAGMENTS):
        return True
    if _looks_like_degenerate_repetition(normalized):
        return True
    if _looks_like_garbled_text(text or ""):
        return True
    return False


def normalize_model_response(
    response_text: Optional[str],
    thinking_text: Optional[str] = None,
) -> tuple[str, str]:
    thinking, response = split_ollama_response_parts(
        response_text=response_text,
        thinking_text=thinking_text,
    )
    if is_invalid_model_output(response):
        logger.warning("Model produced invalid output preview: %r", response[:300])
        return "", _fallback_invalid_generation_response()
    return thinking, response


async def _call_ollama(
    prompt: str,
    options: Optional[dict] = None,
    images: Optional[Sequence[str]] = None,
    thinking: Optional[bool] = None,
) -> dict[str, str]:
    ollama_url = f"{settings.ollama_base_url}/api/generate"
    payload = build_ollama_generate_payload(
        prompt,
        stream=False,
        images=images,
        thinking=thinking,
    )
    if options:
        if "options" in payload:
            payload["options"].update(options)
        else:
            payload["options"] = dict(options)
    async with httpx.AsyncClient(timeout=settings.ollama_timeout) as client:
        response = await client.post(ollama_url, json=payload)
        response.raise_for_status()
        data = response.json()
        thinking, message = normalize_model_response(
            response_text=data.get("response"),
            thinking_text=data.get("thinking"),
        )
        return {"thinking": thinking, "message": message}


def _fallback_response(original_prompt: str) -> str:
    return (
        f"{t('ollama_not_available')} "
        f"{t('ollama_no_response', url=settings.ollama_base_url, model=settings.ollama_model)} "
        f"'{original_prompt}'. "
        f"{t('ollama_check_service')}"
    )


def _fallback_timeout_response(original_prompt: str) -> str:
    timeout_seconds = f"{settings.ollama_timeout:g}"
    return (
        "[Ollama reageert te langzaam] "
        f"Geen antwoord ontvangen binnen {timeout_seconds}s van "
        f"{settings.ollama_base_url} met model '{settings.ollama_model}'. "
        f"Je vroeg: '{original_prompt}'. "
        "Dit gebeurt vaak bij een koude start of het laden van een groot model; probeer het opnieuw."
    )


async def handle_ask(
    req: ChatRequest,
    history: Optional[Sequence[ChatMessage]] = None,
    documents: Optional[Sequence[ParsedDocument]] = None,
    final_prompt_override: Optional[str] = None,
    citations: Optional[List[Dict[str, Any]]] = None,
    web_results: Optional[Sequence[WebSearchResult]] = None,
) -> dict:
    history_to_use = req.history if history is None else history
    current_images = normalize_images(req.images)
    prompt_images = collect_prompt_images(history_to_use, current_images=current_images)
    final_prompt = final_prompt_override or build_augmented_prompt(
        req.prompt,
        history_to_use,
        images_count=len(current_images),
        documents=documents,
        mode=req.mode,
        web_results=web_results,
    )
    log_prompt(final_prompt)
    thinking = ""
    try:
        result = await _call_ollama(final_prompt, images=prompt_images, thinking=req.thinking)
        message = result.get("message", "").strip()
        thinking = result.get("thinking", "").strip()
    except httpx.TimeoutException as exc:  # pragma: no cover - netwerkfout
        logger.warning(
            "Timeout bij Ollama (timeout=%ss, url=%s, model=%s): %r",
            settings.ollama_timeout,
            settings.ollama_base_url,
            settings.ollama_model,
            exc,
        )
        message = _fallback_timeout_response(req.prompt)
    except httpx.RequestError as exc:  # pragma: no cover - netwerkfout
        logger.warning(
            "Kan niet verbinden met Ollama (url=%s, model=%s): %s",
            settings.ollama_base_url,
            settings.ollama_model,
            exc,
        )
        message = _fallback_response(req.prompt)
    except Exception as exc:  # pragma: no cover - netwerkfout
        logger.warning(
            "Onverwachte Ollama fout (url=%s, model=%s): %r",
            settings.ollama_base_url,
            settings.ollama_model,
            exc,
        )
        message = _fallback_response(req.prompt)
    if not message:
        message = t("no_response_generated")
    web_payload: List[Dict[str, Any]] = []
    if web_results:
        web_payload = [item.to_dict() for item in web_results]
    return {
        "message": message,
        "thinking": thinking,
        "citations": list(citations or []),
        "web_results": web_payload,
    }

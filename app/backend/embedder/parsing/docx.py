"""DOCX parser.

Gebruikt :mod:`python-docx`. Splitst op headings (paragraph-style ``Heading N``)
zoals de PhpWord-versie deed met ``Title``-elementen.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .normalizer import normalize_text
from .parsed import ParsedDocument, slugify


_HEADING_STYLE_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)


class DocxParser:
    _SUPPORTED_MIMES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self._SUPPORTED_MIMES

    def parse(self, path: str, options: dict[str, Any]) -> ParsedDocument:
        document = Document(path)
        file_path = Path(path)

        core_props = document.core_properties

        doc_id = options.get("doc_id") or slugify(file_path.stem)
        title = options.get("title") or core_props.title or file_path.stem

        sections = self._extract_sections(document)

        metadata = {
            "author": core_props.author or None,
            "company": getattr(core_props, "company", None) or None,
            "description": core_props.comments or None,
            "keywords": core_props.keywords or None,
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
        }
        metadata = {k: v for k, v in metadata.items() if v}

        return ParsedDocument(
            doc_id=doc_id,
            title=title,
            sections=sections,
            metadata=metadata,
            description=options.get("description") or core_props.comments or None,
            category=options.get("category"),
        )

    def get_mappable_fields(self, path: str) -> dict[str, str]:
        return {}

    def requires_mapping(self) -> bool:
        return False

    # ---- internals ---------------------------------------------------------

    def _extract_sections(self, document) -> list[dict]:
        sections: list[dict] = []
        current_header: dict | None = None
        buffer: list[str] = []
        section_index = 0

        for block in _iter_block_items(document):
            if isinstance(block, Paragraph):
                level = _heading_level(block)
                if level is not None:
                    if current_header is not None and _has_text(buffer):
                        sections.append(
                            self._make_section(section_index, current_header, buffer)
                        )
                        section_index += 1
                    title_text = block.text.strip()
                    current_header = {
                        "title": title_text,
                        "slug": slugify(title_text, fallback=f"sectie-{section_index + 1}"),
                        "level": level,
                    }
                    buffer = []
                else:
                    text = block.text
                    if _is_list_item(block) and text.strip():
                        text = f"- {text.strip()}"
                    buffer.append(text)
            elif isinstance(block, Table):
                buffer.append(self._extract_table_text(block))

        if _has_text(buffer):
            if current_header is None:
                current_header = {"title": "Inhoud", "slug": "inhoud", "level": 0}
            sections.append(self._make_section(section_index, current_header, buffer))

        if not sections:
            joined = "\n".join(_iter_text_all(document))
            sections.append(
                {
                    "title": "Inhoud",
                    "slug": "inhoud",
                    "order_index": 0,
                    "text": normalize_text(joined),
                    "metadata": {},
                }
            )

        return sections

    @staticmethod
    def _make_section(order_index: int, header: dict, buffer: list[str]) -> dict:
        return {
            "title": header["title"],
            "slug": header["slug"],
            "order_index": order_index,
            "text": normalize_text("\n".join(buffer)),
            "metadata": {"level": header.get("level", 0)},
        }

    @staticmethod
    def _extract_table_text(table: Table) -> str:
        lines = ["\n"]
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")
        return "\n".join(lines)


# ---- helpers ----------------------------------------------------------------


def _iter_block_items(document) -> Iterator:
    """Yield Paragraph and Table objects in document order.

    python-docx exposes ``document.paragraphs`` en ``document.tables`` apart,
    waardoor de oorspronkelijke volgorde tussen tables en paragraphs verloren
    gaat. We iteren daarom rechtstreeks over de XML-body en wikkelen elk
    element in het juiste objecttype.
    """

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _heading_level(paragraph: Paragraph) -> int | None:
    style = getattr(paragraph.style, "name", "") or ""
    if not style:
        return None
    match = _HEADING_STYLE_RE.match(style.strip())
    if match:
        return int(match.group(1))
    if style.strip().lower() == "title":
        return 0
    return None


def _is_list_item(paragraph: Paragraph) -> bool:
    style = getattr(paragraph.style, "name", "") or ""
    return "list" in style.lower()


def _has_text(buffer: list[str]) -> bool:
    return any(part.strip() for part in buffer)


def _iter_text_all(document) -> Iterator[str]:
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            yield block.text
        elif isinstance(block, Table):
            for row in block.rows:
                yield " | ".join(cell.text.strip() for cell in row.cells)
